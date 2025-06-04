import re
import os
import time
import json
from typing import List
from dotenv import load_dotenv

import win32com.client
from docx import Document
from docx2pdf import convert

from openai import OpenAI
from fastapi import FastAPI, HTTPException, Query
from fastapi.responses import JSONResponse
import uvicorn
from fastapi.middleware.cors import CORSMiddleware

# Load environment variables from .env file for API keys, etc.
load_dotenv()
api_key = os.getenv("API_KEY")

# Initialize OpenAI client (not used in main flow, but available)
client = OpenAI(api_key=api_key)

def convert_pdf_to_word(pdf_path, docx_path):
    """
    Convert a PDF file to a Word (.docx) file using Microsoft Word automation.
    """
    print("Converting PDF to Word using Microsoft Word...")
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(pdf_path)
    doc.SaveAs(docx_path, FileFormat=16)
    doc.Close()
    word.Quit()
    print("Conversion done.")

def tokenize_with_punctuation(text):
    """
    Tokenizer that splits punctuation as separate tokens.
    """
    return re.findall(r"\w+|[^\w\s]", text, re.UNICODE)

def convert_word_to_pdf(updated_docx_path, final_pdf_path):
    """
    Convert a Word (.docx) file back to PDF.
    """
    print("Converting back to PDF...")
    convert(updated_docx_path, final_pdf_path)
    print(f"Final PDF saved as: {final_pdf_path}")

def gpt_proofread(text):
    """
    Use OpenAI GPT to proofread and classify changes in a sentence.
    Returns a JSON object with original/corrected text, tokens, and changes.
    """
    system_msg = (
        "You are a grammar corrector that also classifies changes and suggests synonyms in a formal manner. "
        "You will receive a sentence to correct. Return a JSON object with the following keys:\n"
        "- 'original': the original input sentence\n"
        "- 'corrected': the corrected version of the sentence\n"
        "- 'original_token': list of objects with 'idx' and 'word' from the original sentence\n"
        "- 'proofread_token': list of objects with 'idx' and 'word' from the corrected sentence\n"
        "- 'changes': a list of changes, where each item includes:\n"
        "  - 'type': one of 'replaced', 'inserted', 'removed', or 'corrected'\n"
        "  - 'original_idx': index in original_token (can be null for insertions)\n"
        "  - 'proofread_idx': index in proofread_token (can be null for removals)\n"
        "  - 'original_word': word or punctuation in the original sentence\n"
        "  - 'proofread_word': corrected word or punctuation\n"
        "  - 'suggestion': up to 3 synonyms (only for 'replaced' or 'inserted')\n"
        "Important: Treat punctuation changes (e.g., commas, periods) as valid changes and reflect them in the tokens and change list."
    )

    user_msg = f"Original sentence:\n{text}\n\nPlease return only the JSON."

    response = client.chat.completions.create(
        # model="gpt-4o-mini",
        model="gpt-4.1",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg}
        ],
        temperature=0.3,
    )

    content = response.choices[0].message.content.strip()

    try:
        match = re.search(r'\{.*\}', content, re.DOTALL)
        if not match:
            raise ValueError("No JSON object found in GPT response")

        json_str = match.group(0)
        # Remove trailing commas before closing brackets/braces
        json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)
        result = json.loads(json_str)

        if "corrected" not in result or "changes" not in result:
            raise ValueError("Missing expected keys in GPT response")

        return result
    except Exception as e:
        print(f"Failed to parse GPT response for input: {text}\nError: {e}")
        raise ValueError(f"Invalid GPT JSON response: {content}")


def correct_paragraphs(docx_path, updated_docx_path, json_output_path, pdf_id="example_pdf_001"):
    """
    Correct grammar in each paragraph of a Word document using GPT,
    update the document, and save a JSON report of all corrections.
    Returns the total number of improved paragraphs.
    """
    print(f"Opening document: {docx_path}")
    doc = Document(docx_path)
    data = {"pdf_id": pdf_id, "paragraphs": []}
    total_improvements = 0
    
    # Regex pattern for matching both Arabic numerals and Roman numerals
    # Arabic numerals: matches digits like 1, 2, 10, 123
    # Roman numerals: matches Roman numeral patterns like I, IV, V, IX, X, XII, etc.
    numeral_pattern = r"^(?:[IVXLCDM]+|\d+)\s"

    print("Starting proofreading of document paragraphs...")
    for idx, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        if not original_text.strip():
            print(f"Paragraph {idx + 1} is empty. Skipping.")
            continue

        print(f"Processing paragraph {idx + 1}:")
        print(f"Original: {original_text}")
        
        # Check if the paragraph starts with a number or a Roman numeral followed by space
        match = re.match(numeral_pattern, original_text)
        if match:
            # Separate the numeral (Arabic or Roman) from the rest of the paragraph
            numeral_part = match.group(0)
            text_part = original_text[len(numeral_part):]
        else:
            numeral_part = ""
            text_part = original_text

        try:
            gpt_response = gpt_proofread(text_part)
        except Exception as e:
            print(f"Error processing paragraph {idx + 1}: {e}")
            continue

        corrected_text_part = gpt_response.get("corrected", text_part)
        corrected_text = numeral_part + corrected_text_part
        if corrected_text != original_text:
            print(f"Paragraph {idx + 1} corrected.")
            total_improvements += 1
        else:
            print(f"No changes made to paragraph {idx + 1}.")

        para_id = idx + 1
        data["paragraphs"].append({
            "paragraph_id": para_id,
            "original": gpt_response.get("original"),
            "proofread": gpt_response.get("corrected"),
            "original_token": gpt_response.get("original_token", []),
            "proofread_token": gpt_response.get("proofread_token", []),
            "original_text": [
                {
                    "index": ch.get("original_idx"),
                    "word": ch.get("original_word"),
                    "type": "error"
                }
                for ch in gpt_response.get("changes", [])
                if ch.get("type") != "inserted" and ch.get("original_idx") is not None
            ],
            "revised_text": [
                {
                    "index": ch.get("proofread_idx"),
                    "word": ch.get("proofread_word"),
                    "type": ch.get("type"),
                    "suggestions": ch.get("suggestion", [ch.get("proofread_word")])
                }
                for ch in gpt_response.get("changes", [])
                if ch.get("proofread_idx") is not None
            ]
        })

        # Update Word paragraph while preserving formatting
        if paragraph.runs:
            print(f"Preserving formatting for paragraph {idx + 1}")
            ref_run = paragraph.runs[0]
            paragraph.clear()
            new_run = paragraph.add_run(corrected_text)
            new_run.font.name = ref_run.font.name
            new_run.bold = ref_run.bold
            new_run.italic = ref_run.italic
            new_run.underline = ref_run.underline
            new_run.font.size = ref_run.font.size
            if ref_run.font.color and ref_run.font.color.rgb:
                new_run.font.color.rgb = ref_run.font.color.rgb
        else:
            paragraph.text = corrected_text

    print(f"Saving updated document to: {updated_docx_path}")
    doc.save(updated_docx_path)
    print("Document saved successfully.")

    print(f"Saving JSON report to: {json_output_path}")
    with open(json_output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print("JSON report saved successfully.")

    print(f"Total improved paragraphs: {total_improvements}")
    return total_improvements


def update_changes_on_pdf(final_pdf_path, updated_docx_path, json_output_path, paragraph_id):
    """
    Updates the paragraphs in a Word document based on proofread data from a JSON file.
    Only updates paragraphs whose IDs are in paragraph_id.
    Returns the number of paragraphs updated.
    """
    
    print("🔍 Loading JSON proofreading results...")
    with open(json_output_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    paragraphs_data = data.get("paragraphs", [])
    paragraph_id_set = set(str(pid) for pid in paragraph_id)  # ensure all are strings
    print(f"📋 Paragraph IDs to update: {paragraph_id_set}")
    
    print(f"📄 Loading DOCX: {updated_docx_path}")
    doc = Document(updated_docx_path)
    updated_count = 0

    for para in paragraphs_data:
        pid = str(para.get("paragraph_id"))
        proofread_text = para.get("proofread")

        if pid in paragraph_id_set:
            para_index = int(pid) - 1  # Adjust for 0-based indexing
            print(f"\n🛠 Updating Paragraph ID: {pid} (Index: {para_index})")
            print(f"✅ New Proofread Text: {repr(proofread_text)}")

            if 0 <= para_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[para_index]
                print(f"✏️ Original Text: {repr(paragraph.text)}")

                if paragraph.runs:
                    ref_run = paragraph.runs[0]
                    paragraph.clear()
                    new_run = paragraph.add_run(proofread_text)
                    new_run.font.name = ref_run.font.name
                    new_run.bold = ref_run.bold
                    new_run.italic = ref_run.italic
                    new_run.underline = ref_run.underline
                    new_run.font.size = ref_run.font.size
                    if ref_run.font.color and ref_run.font.color.rgb:
                        new_run.font.color.rgb = ref_run.font.color.rgb
                else:
                    paragraph.text = proofread_text

                print("✅ Paragraph updated successfully.")
                updated_count += 1
            else:
                print(f"❌ Invalid paragraph index: {para_index}")
        else:
            print(f"⏭️ Skipping paragraph ID {pid} (not in list)")

    print(f"\n📁 Saving updated DOCX: {updated_docx_path}")
    doc.save(updated_docx_path)

    print(f"\n✅ Total paragraphs updated: {updated_count}")
    return updated_count
# Initialize FastAPI app
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)

@app.get("/api/grammar-check")
async def main(
    mode: int = Query(...),
    file_code: str = Query(...),
    paragraph_id: str = Query(...)
):
    # Clean input like "[1,2,3]" or "1,2,3"
    cleaned = re.sub(r"[\[\]\s]", "", paragraph_id)
    paragraph_id_list = cleaned.split(",")
    """
    Main API endpoint for grammar checking and PDF processing.
    - mode="0": Full process (PDF→Word→Proofread→PDF)
    - mode="1": Update only selected paragraphs (using paragraph_id)
    Returns output filenames, total improvements, and elapsed time.
    """
    start_time = time.time()

    # Build file paths based on file_code
    pdf_path = os.path.abspath(f"original_pdfs/{file_code}.pdf")
    docx_path = os.path.abspath(f"parsing_words/{file_code}_temp.docx")
    updated_docx_path = os.path.abspath(f"parsing_words/{file_code}_updated.docx")
    final_pdf_path = os.path.abspath(f"processed_pdfs/{file_code}.pdf")
    json_output_path = os.path.abspath(f"jsons/{file_code}.json")
    
    print("Debug Paths:")
    print("PDF Path:             ", os.path.abspath(f"original_pdfs/{file_code}.pdf"))
    print("Temp DOCX Path:       ", os.path.abspath(f"parsing_words/{file_code}_temp.docx"))
    print("Updated DOCX Path:    ", os.path.abspath(f"parsing_words/{file_code}_updated.docx"))
    print("Final PDF Path:       ", os.path.abspath(f"processed_pdfs/{file_code}.pdf"))
    print("JSON Output Path:     ", os.path.abspath(f"jsons/{file_code}.json"))

    # Main processing logic based on mode
    if mode == 0:
        # Full process: convert, proofread, and save all
        convert_pdf_to_word(pdf_path, docx_path)
        total_improvements = correct_paragraphs(docx_path, updated_docx_path, json_output_path)
        convert_word_to_pdf(updated_docx_path, final_pdf_path)
    else:
        # Only update selected paragraphs
        convert_pdf_to_word(final_pdf_path, docx_path)
        total_improvements = update_changes_on_pdf(final_pdf_path, updated_docx_path, json_output_path, paragraph_id)
        convert_word_to_pdf(updated_docx_path, final_pdf_path)

    elapsed = time.time() - start_time
    print(f"Total processing time: {elapsed:.2f} seconds")
    print(f"Total Error Found: {total_improvements}")

    # Return summary as JSON
    return {
        "json_filename": os.path.basename(json_output_path),
        "final_pdf_filename": os.path.basename(final_pdf_path),
        "total_improvements": total_improvements,
        "elapsed_time_seconds": round(elapsed, 2)
    }
    
# @

if __name__ == "__main__":
    # Run the FastAPI app with Uvicorn
    uvicorn.run(app, host="0.0.0.0", port=5000)